from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "output" / "docx"

PORTFOLIO = "https://wsmykm3.github.io/ProfolioWSM/"
LINKEDIN = "https://www.linkedin.com/in/siming-wang-321a18303/"
GITHUB = "https://github.com/WSMYKM3"
EMAIL = "simingvv@gmail.com"

PROJECT_URLS = {
    "Reroll": "https://wsmykm3.github.io/ProfolioWSM/projects/post-7",
    "I AND AI: MIRROR": "https://wsmykm3.github.io/ProfolioWSM/projects/post-3",
    "Signie": "https://wsmykm3.github.io/ProfolioWSM/projects/post-2",
    "Sorting Factory": "https://wsmykm3.github.io/ProfolioWSM/projects/post-8",
    "The Tool Box": "https://wsmykm3.github.io/ProfolioWSM/projects/post-5",
}

# compact_reference_guide, with a named resume_a4_single_column override.
# User-directed overrides: A4, 14 mm side margins, 10 pt body, one-column,
# 1.03-1.05 line spacing, no tables, no running furniture on one-page files.
INK = RGBColor(20, 33, 43)
ACCENT = RGBColor(31, 93, 131)
MUTED = RGBColor(78, 94, 104)
LINK = RGBColor(22, 93, 145)
FONT = "Arial"
CJK_FONT = "Arial Unicode MS"


def set_rfonts(rpr, ascii_font=FONT, east_asia=CJK_FONT):
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:cs"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia)


def style_font(style, size, color=INK, bold=False, italic=False):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    set_rfonts(rpr)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")


def set_run(run, size=10, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    set_rfonts(rpr)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")


def add_hyperlink(paragraph, text, url, size=9.2, color=LINK, bold=False, underline=False):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    set_rfonts(rpr)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(color))
    rpr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(round(size * 2))))
    rpr.append(size_el)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(round(size * 2))))
    rpr.append(size_cs)
    if bold:
        rpr.append(OxmlElement("w:b"))
    underline_el = OxmlElement("w:u")
    underline_el.set(qn("w:val"), "single" if underline else "none")
    rpr.append(underline_el)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")
    rpr.append(lang)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def make_bullet_numbering(doc):
    root = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    set_rfonts(rpr)
    lvl.append(rpr)
    abstract.append(lvl)
    root.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)


def add_style(doc, name, size, color=INK, bold=False, italic=False, before=0, after=0, line=1.0,
              keep_next=False, keep_together=False):
    styles = doc.styles
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style_font(style, size, color, bold, italic)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together
    return style


def configure_document(title, subject, keywords, master=False):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(14 if not master else 15)
    sec.right_margin = Mm(14 if not master else 15)
    sec.top_margin = Mm(10 if not master else 12)
    sec.bottom_margin = Mm(10 if not master else 12)
    sec.header_distance = Mm(5)
    sec.footer_distance = Mm(5)

    normal = doc.styles["Normal"]
    style_font(normal, 10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.4)
    normal.paragraph_format.line_spacing = 1.05

    heading = doc.styles["Heading 1"]
    style_font(heading, 10.6 if not master else 11.0, ACCENT, True)
    heading.paragraph_format.space_before = Pt(4.6)
    heading.paragraph_format.space_after = Pt(1.6)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    add_style(doc, "Resume Name", 21, INK, True, after=0, line=1.0, keep_next=True)
    add_style(doc, "Resume Headline", 10.8, ACCENT, True, after=0.5, line=1.0, keep_next=True)
    add_style(doc, "Resume Contact", 9.1, MUTED, after=0, line=1.0, keep_next=True)
    add_style(doc, "Resume Summary", 10, INK, after=1.5, line=1.07, keep_together=True)
    add_style(doc, "Resume Item Header", 10, INK, True, before=1.8, after=0.4, line=1.0, keep_next=True)
    add_style(doc, "Resume Compact", 10, INK, after=0.7, line=1.03, keep_together=True)
    add_style(doc, "Resume Small", 9.2, MUTED, after=0.5, line=1.0, keep_together=True)

    props = doc.core_properties
    props.author = "Siming Wang"
    props.last_modified_by = "Siming Wang"
    props.title = title
    props.subject = subject
    props.keywords = keywords
    props.comments = "ATS-compatible, single-column resume"

    settings = doc.settings.element
    theme_lang = settings.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_lang)
    theme_lang.set(qn("w:val"), "en-US")
    theme_lang.set(qn("w:eastAsia"), "zh-CN")

    if master:
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        run = fp.add_run("Siming Wang | Master Resume")
        set_run(run, 8.5, MUTED)

    return doc, make_bullet_numbering(doc)


def add_header(doc, headline_lines):
    p = doc.add_paragraph(style="Resume Name")
    set_run(p.add_run("SIMING WANG"), 21, INK, True)
    for line in headline_lines:
        p = doc.add_paragraph(style="Resume Headline")
        set_run(p.add_run(line), 10.8, ACCENT, True)

    p = doc.add_paragraph(style="Resume Contact")
    set_run(p.add_run("Shanghai, China | "), 9.1, MUTED)
    add_hyperlink(p, EMAIL, f"mailto:{EMAIL}", 9.1, LINK)
    set_run(p.add_run(" | Portfolio: "), 9.1, MUTED)
    add_hyperlink(p, "wsmykm3.github.io/ProfolioWSM", PORTFOLIO, 9.1, LINK)
    p = doc.add_paragraph(style="Resume Contact")
    add_hyperlink(p, "linkedin.com/in/siming-wang-321a18303", LINKEDIN, 9.1, LINK)
    set_run(p.add_run(" | "), 9.1, MUTED)
    add_hyperlink(p, "github.com/WSMYKM3", GITHUB, 9.1, LINK)


def add_section(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    set_run(p.add_run(title), 10.6, ACCENT, True)
    return p


def add_summary(doc, text):
    p = doc.add_paragraph(style="Resume Summary")
    set_run(p.add_run(text), 10, INK)


def add_item_header(doc, title, meta, url=None):
    p = doc.add_paragraph(style="Resume Item Header")
    if url:
        add_hyperlink(p, title, url, 10, ACCENT, True)
    else:
        set_run(p.add_run(title), 10, INK, True)
    set_run(p.add_run(f" | {meta}"), 9.6, MUTED, True)
    return p


def add_bullet(doc, num_id, text):
    p = doc.add_paragraph(style="Resume Compact")
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(0.8)
    p.paragraph_format.line_spacing = 1.03
    p.paragraph_format.keep_together = True
    set_run(p.add_run(text), 10, INK)
    return p


def add_label_line(doc, label, value, size=10):
    p = doc.add_paragraph(style="Resume Compact")
    set_run(p.add_run(label), size, INK, True)
    set_run(p.add_run(value), size, INK)
    return p


def add_project(doc, num_id, name, meta, bullets):
    add_item_header(doc, name, meta, PROJECT_URLS[name])
    for bullet in bullets:
        add_bullet(doc, num_id, bullet)


def add_experience(doc, num_id, title, meta, bullets):
    add_item_header(doc, title, meta)
    for bullet in bullets:
        add_bullet(doc, num_id, bullet)


def build_ct_en():
    doc, num = configure_document(
        "Siming Wang - Creative Technologist",
        "Creative Technologist resume for Shanghai roles",
        "Creative Technologist, Generative AI, Interactive Experiences, Rapid Prototyping, Unity, Unreal Engine, TouchDesigner, ARKit, Python",
    )
    add_header(doc, ["CREATIVE TECHNOLOGIST | GENERATIVE AI | INTERACTIVE EXPERIENCES | RAPID PROTOTYPING"])
    add_section(doc, "SUMMARY")
    add_summary(doc, "Creative Technologist who bridges creative direction and engineering to turn briefs into functional AI-powered prototypes and interactive experiences. Combines generative AI workflows, computer vision, spatial computing and real-time engines from concept validation through system integration and live exhibition.")

    add_section(doc, "SELECTED PROJECTS")
    add_project(doc, num, "Reroll", "Independent AI Filmmaking + AR Prototype | Apr-Aug 2026", [
        "Built a functional pipeline that converts a reference image into editable scene data with Python, SAM 3 and DPT, then brings it into an iPhone director tool built with SwiftUI, ARKit and RealityKit.",
        "Enabled object-level edits, physical camera-path capture and Apple Speech notes; an AI agent synthesizes scene, camera and voice inputs into a refined prompt for video generation.",
    ])
    add_project(doc, num, "I AND AI: MIRROR", "Interactive AI Installation | 2025", [
        "Delivered the real-time interaction system for an Immersive Arts UK-supported exhibition at Inspace, Edinburgh: TouchDesigner/Python state logic, Vosk wake-word activation, OSC and Unreal Engine MetaHuman speech/lip sync.",
        "Integrated tablet triggers, inactivity fallbacks and multi-device communication for a three-day live run with 422 interactions, including 287 complete experiences.",
    ])
    add_project(doc, num, "Signie", "XR ASL Learning + Translation Prototype | Mar-Jun 2025", [
        "Developed a Unity 6/Meta Quest prototype with hand-tracked learning interactions, guided tutor animation, micro-gesture control and a Wit.ai-driven animation state machine; presented at AWE USA 2025.",
    ])

    add_section(doc, "PROFESSIONAL EXPERIENCE")
    add_experience(doc, num, "XR Engineer, Part-time", "TeknTrash Robotics | Remote, UK | Mar 2026-Present", [
        "Develop XR applications and integrate XR hardware with robotic platforms for real-time control and human-robot interaction.",
        "Support XR-derived data capture and system experiments for Vision-Language-Action workflows; document interfaces and test procedures with cross-functional collaborators.",
    ])
    add_experience(doc, num, "Technical Artist / Creative Technologist, Freelance", "6Liè Projects | Remote, UK | Mar 2025-Jan 2026", [
        "Built real-time modules for I AND AI: MIRROR, connecting TouchDesigner, Python/Vosk and Unreal Engine/MetaHuman through OSC for a live immersive installation.",
    ])
    add_experience(doc, num, "Visual Artist / Creative Technologist Intern", "Shanghai Chaomo Studio | Shanghai | Sep 2021-Jul 2022", [
        "Built JavaScript/Python interactive and generative prototypes and visual systems for thematic exhibitions and installations.",
    ])

    add_section(doc, "CAPABILITIES")
    add_label_line(doc, "AI & Creative Prototyping: ", "Generative AI workflows, AI prompt synthesis, computer vision, Python, SAM 3, DPT, Apple Speech, Wit.ai, Vosk")
    add_label_line(doc, "Real-Time & Spatial: ", "Unity 6, C#, Unreal Engine, Blueprints, SwiftUI, ARKit, RealityKit, Meta Quest")
    add_label_line(doc, "Interactive Systems: ", "TouchDesigner, OSC, JavaScript, hand tracking, micro-gestures, state machines, multi-device integration")
    add_label_line(doc, "3D & Production: ", "Blender, MetaHuman, MotionBuilder, motion capture, real-time animation, exhibition integration")

    add_section(doc, "EDUCATION | TALKS & EXHIBITIONS")
    add_label_line(doc, "MFA Computational Arts | ", "Goldsmiths, University of London | 2022-2024")
    add_label_line(doc, "BA Visual Communication Design | ", "East China University of Science and Technology | 2019-2022")
    add_label_line(doc, "AWE USA 2025 Speaker | ", "Presented Signie")
    add_label_line(doc, "I AND AI: MIRROR | ", "Immersive Arts UK-supported pop-up exhibition and performance, Inspace, Edinburgh | 2025")
    return doc


def build_ct_cn():
    doc, num = configure_document(
        "Siming Wang - 创意技术 / Creative Technologist",
        "面向上海创意技术与 AI 创意产品岗位的中文简历",
        "创意技术, Creative Technologist, Generative AI, 互动体验, Rapid Prototyping, Unity, TouchDesigner, ARKit, Python",
    )
    add_header(doc, ["创意技术 / CREATIVE TECHNOLOGIST", "GENERATIVE AI | INTERACTIVE EXPERIENCES | RAPID PROTOTYPING"])
    add_section(doc, "个人简介 / PROFILE")
    add_summary(doc, "面向 AI 创意产品、把创意 brief 转化为可运行体验的 Creative Technologist，横跨 Generative AI、实时交互、XR 与计算机视觉。能够从概念、体验流程和视觉语言出发，独立完成 Rapid Prototyping 与现场系统集成。")

    add_section(doc, "精选项目 / SELECTED PROJECTS")
    add_project(doc, num, "Reroll", "个人 AI 影像 + AR 功能原型 | 2026.04-08", [
        "面向不熟悉 3D 工具的创作者，将参考图转为可编辑的 iPhone AR 场景，可调整对象、构图与运镜，并录入绑定到对象的语音导演指令。",
        "独立打通 Python、SAM 3/DPT、SwiftUI、ARKit/RealityKit、Apple Speech 与 AI prompt synthesis，输出供 AI 视频生成使用的精炼提示词。",
    ])
    add_project(doc, num, "I AND AI: MIRROR", "沉浸式 AI 互动装置 | 2025", [
        "负责 Unreal Engine/MetaHuman、TouchDesigner、Python/Vosk、OSC、唤醒词、交互状态机与多设备联调，将人类与 AI 亲密关系的概念转为现场数字化身体验。",
        "项目获 Immersive Arts UK 支持并在 Inspace Edinburgh 展演，3 天记录 422 次参与、287 次完整体验。",
    ])
    add_project(doc, num, "Signie", "XR ASL Learning + Translation Prototype | 2025.03-06", [
        "在 Unity 6/Meta Quest 中构建 hand tracking、micro-gestures、引导动画与状态管理，以 Wit.ai Speech-to-Text 驱动手语动画；原型于 AWE USA 2025 展示。",
    ])

    add_section(doc, "工作经历 / EXPERIENCE")
    add_experience(doc, num, "XR Engineer（Part-time）", "TeknTrash Robotics | Remote UK | 2026.03-至今", [
        "围绕机器人系统搭建 XR 数据采集、实时控制与硬件集成原型，并维护接口文档、测试流程和系统行为说明。",
    ])
    add_experience(doc, num, "Technical Artist / Creative Technologist（Freelance）", "6Liè Projects | Remote UK | 2025.03-2026.01", [
        "交付沉浸式项目的实时互动与 AI 集成模块；在 MIRROR 中负责 OSC 通信、TouchDesigner 状态机、语音唤醒与现场联调。",
    ])
    add_experience(doc, num, "Visual Artist / Creative Technologist（Intern）", "Shanghai Chaomo Studio | 上海 | 2021.09-2022.07", [
        "使用 JavaScript/Python 制作互动与生成式视觉原型，为主题展览和装置设计实时视觉系统。",
    ])

    add_section(doc, "专业能力 / CAPABILITIES")
    add_label_line(doc, "AI & Creative Prototyping：", "Python、Generative AI workflow、AI prompt synthesis、Computer Vision、SAM 3、DPT")
    add_label_line(doc, "Real-Time & Spatial：", "Unity 6/C#、Unreal Engine/Blueprints、ARKit、RealityKit、Meta Quest")
    add_label_line(doc, "Interactive Systems：", "TouchDesigner、OSC、Apple Speech、Vosk、Wit.ai、hand tracking、micro-gestures")
    add_label_line(doc, "3D & Production：", "Blender、MetaHuman、MotionBuilder、Figma、Photoshop、Premiere Pro")

    add_section(doc, "教育 | 演讲与展览")
    add_label_line(doc, "MFA Computational Arts | ", "Goldsmiths, University of London | 2022-2024")
    add_label_line(doc, "BA Visual Communication Design | ", "East China University of Science and Technology | 2019-2022")
    add_label_line(doc, "AWE USA 2025 Speaker | ", "Signie")
    add_label_line(doc, "I AND AI: MIRROR | ", "Immersive Arts UK-supported Pop-up Exhibition | Inspace, Edinburgh | 2025")
    return doc


def build_xr_en():
    doc, num = configure_document(
        "Siming Wang - XR Engineer",
        "XR Engineer resume for spatial computing and human-robot interaction roles",
        "XR Engineer, Spatial Computing, Computer Vision, Human-Robot Interaction, Unity, C#, Python, ARKit, RealityKit, FastAPI, WebSocket",
    )
    add_header(doc, ["XR ENGINEER | SPATIAL COMPUTING | COMPUTER VISION | HUMAN-ROBOT INTERACTION"])
    add_section(doc, "SUMMARY")
    add_summary(doc, "XR engineer and creative technologist building real-time systems across spatial computing, robotics and computer vision. Develops end-to-end prototypes with Unity/C#, Python, ARKit/RealityKit and FastAPI/WebSocket, translating perception and AI outputs into responsive human-robot and embodied interactions.")

    add_section(doc, "PROFESSIONAL EXPERIENCE")
    add_experience(doc, num, "XR Engineer, Part-time", "TeknTrash Robotics | Remote, UK | Mar 2026-Present", [
        "Develop XR applications and integrate XR hardware with robotic platforms to prototype real-time control and human-robot interaction.",
        "Contribute XR-derived data and system-integration work to Vision-Language-Action experiments; document interfaces, test procedures and system behavior for cross-functional teams.",
    ])
    add_experience(doc, num, "Technical Artist / Creative Technologist, Freelance", "6Liè Projects | Remote, UK | Mar 2025-Jan 2026", [
        "Delivered a TouchDesigner state machine, Python/Vosk wake-word activation, OSC-to-Unreal MetaHuman connection and iPad trigger/fallback for a live immersive installation.",
    ])
    add_experience(doc, num, "Visual Artist / Creative Technologist Intern", "Shanghai Chaomo Studio | Shanghai | Sep 2021-Jul 2022", [
        "Built interactive and generative prototypes with JavaScript/Python and designed visual systems for exhibitions and installations.",
    ])

    add_section(doc, "SELECTED PROJECTS")
    add_project(doc, num, "Sorting Factory", "AI & Robotics Engineer | Personal Simulation Prototype | Jul 2026", [
        "Built a Unity 6/C# simulation with three parallel SO-101 digital twins, time-aware pick decisions and structured success, failure and abandoned-attempt logging.",
        "Streamed camera frames and ROI metadata to a Python YOLO/ByteTrack service over WebSocket; implemented a FastAPI control room, real-time telemetry and CSV episodes for future robotics training.",
    ])
    add_project(doc, num, "Reroll", "Independent AI + AR Functional Prototype | Apr-Aug 2026", [
        "Built a Python SAM 3/DPT pipeline and a SwiftUI/ARKit/RealityKit iPhone directing tool for spatial edits, camera-path capture, object-linked Apple Speech notes and AI prompt synthesis.",
    ])
    add_project(doc, num, "Signie", "XR ASL Prototype Presented at AWE USA 2025 | Mar-Jun 2025", [
        "Developed hand-tracked Unity 6/Meta Quest learning flows, micro-gesture control and Wit.ai-driven sign animation through an application state manager.",
    ])

    add_section(doc, "TECHNICAL SKILLS")
    add_label_line(doc, "XR & Real-Time: ", "Unity 6, C#, XR Interaction Toolkit, ARKit, RealityKit, Meta Quest, Unreal Engine, Blueprints, TouchDesigner, OSC")
    add_label_line(doc, "AI, Vision & Robotics: ", "Python, YOLO, ByteTrack, SAM 3, DPT, SO-101, IK, VLA data workflows, real-time telemetry")
    add_label_line(doc, "Services & Interaction: ", "FastAPI, WebSocket, REST API, Apple Speech, Wit.ai, Vosk, hand tracking, micro-gestures")

    add_section(doc, "EDUCATION | TALKS & EXHIBITIONS")
    add_label_line(doc, "MFA Computational Arts | ", "Goldsmiths, University of London | 2022-2024")
    add_label_line(doc, "BA Visual Communication Design | ", "East China University of Science and Technology | 2019-2022")
    add_label_line(doc, "AWE USA 2025 Speaker | ", "Presented Signie")
    add_label_line(doc, "I AND AI: MIRROR | ", "Immersive Arts UK-supported exhibition, Inspace, Edinburgh | 2025")
    return doc


def build_master():
    doc, num = configure_document(
        "Siming Wang - Master Resume",
        "Internal two-page master resume covering Creative Technology and XR Engineering",
        "Creative Technologist, XR Engineer, Generative AI, Spatial Computing, Interactive Systems, Unity, Python, TouchDesigner, ARKit, Robotics",
        master=True,
    )
    add_header(doc, ["CREATIVE TECHNOLOGIST & XR ENGINEER", "GENERATIVE AI | INTERACTIVE EXPERIENCES | SPATIAL COMPUTING | REAL-TIME SYSTEMS"])
    add_section(doc, "PROFILE")
    add_summary(doc, "Creative technologist and XR engineer with an MFA in Computational Arts, building AI-enabled interactive experiences, spatial interfaces and real-time robotics prototypes. Combines visual storytelling, rapid prototyping and hands-on engineering to translate creative briefs and system problems into functional prototypes, exhibited installations and testable workflows.")

    add_section(doc, "PROFESSIONAL EXPERIENCE")
    add_experience(doc, num, "XR Engineer, Part-time", "TeknTrash Robotics | Remote, UK | Mar 2026-Present", [
        "Develop XR applications and integrate XR hardware with robotic platforms to prototype real-time control and human-robot interaction.",
        "Contribute XR-derived data and system-integration work to Vision-Language-Action experiments; document interfaces, test procedures and system behavior for cross-functional teams.",
    ])
    add_experience(doc, num, "Technical Artist / Creative Technologist, Freelance", "6Liè Projects | Remote, UK | Mar 2025-Jan 2026", [
        "Delivered the real-time interaction layer for I AND AI: MIRROR, connecting a TouchDesigner state machine to Unreal Engine MetaHuman through OSC and adding Python/Vosk wake-word control.",
        "Built an iPad interaction trigger and inactivity-fallback interface for the live installation; contributed to work supported by Immersive Arts UK and the UKRI Innovate UK Immersive Tech Network.",
    ])
    add_experience(doc, num, "Visual Artist / Creative Technologist Intern", "Shanghai Chaomo Studio | Shanghai | Sep 2021-Jul 2022", [
        "Built JavaScript/Python interactive and generative prototypes and visual systems for thematic exhibitions and installations.",
    ])

    add_section(doc, "SELECTED PROJECTS")
    add_project(doc, num, "Sorting Factory", "AI & Robotics Engineer | Personal Simulation Prototype | Jul 2026", [
        "Built a Unity 6/C# physical-simulation workflow with three parallel SO-101 digital twins, conveyor workstations, local robot controllers and time-aware pick decisions.",
        "Implemented a Python YOLO/ByteTrack service over WebSocket, a FastAPI browser control room, real-time telemetry and CSV records containing joint data, actions, Track IDs, timing, outcomes and failure reasons.",
    ])
    add_project(doc, num, "Reroll", "Creative Technologist / AI & AR Developer | Independent Prototype | Apr-Aug 2026", [
        "Created a phone-first alternative to traditional 3D previsualization: users start with a reference image, edit the scene in AR, direct the camera physically and add object-linked spoken notes.",
        "Built the Python SAM 3/DPT vision pipeline and iPhone interface with SwiftUI, ARKit, RealityKit, Apple Speech and AI prompt synthesis.",
    ])

    doc.add_page_break()
    add_section(doc, "SELECTED PROJECTS (CONTINUED)")
    add_project(doc, num, "I AND AI: MIRROR", "Real-Time Systems / Technical Artist | Immersive Prototype | Apr-Oct 2025", [
        "Developed the real-time interaction system using Unreal Engine MetaHuman, TouchDesigner, Python and Blender.",
        "Built the chat state machine, Python/Vosk wake-word activation, OSC-driven speech/lip sync, multi-device communication and iPad trigger with inactivity fallback.",
        "Presented at Inspace, Edinburgh as an Immersive Arts UK-supported pop-up exhibition and performance; recorded 422 interactions and 287 complete experiences over three days.",
    ])
    add_project(doc, num, "Signie", "XR Developer | Prototype Presented at AWE USA 2025 | Mar-Jun 2025", [
        "Developed Unity 6 ASL teaching and game-based testing flows with guided animation, hand tracking, micro-gestures and an application state manager.",
        "Connected Wit.ai speech-to-text to an animation state machine for sign-language responses; presented the prototype at AWE USA 2025.",
    ])
    add_project(doc, num, "The Tool Box", "Team Lead / XR Developer | 5-Person Team | 3-Day Prototype, XRCC Berlin 2025", [
        "Led a five-person team and developed a Unity 6 XR product-exploration prototype for Strauss tools, with virtual guides, usage scenarios and incompatible-drill-bit safety warnings.",
        "Prototyped an AI shopping assistant using speech-to-text, the OpenAI API and text-to-speech for safety guidance, prices and product links.",
    ])

    add_section(doc, "CAPABILITIES & SKILLS")
    add_label_line(doc, "AI & Creative Prototyping: ", "Python, JavaScript, OpenAI API, AI prompt synthesis, generative AI workflows, YOLO, ByteTrack, SAM 3, DPT, VLA experimentation")
    add_label_line(doc, "Real-Time & Spatial: ", "Unity 6, C#, XR Interaction Toolkit, ARKit, RealityKit, Meta Quest, Unreal Engine, Blueprints, MetaHuman, SO-101 simulation, IK")
    add_label_line(doc, "Interactive Systems: ", "TouchDesigner, OSC, FastAPI, WebSocket, REST API, Apple Speech, Wit.ai, Vosk, hand tracking, micro-gestures, real-time telemetry")
    add_label_line(doc, "3D & Production: ", "Blender, MotionBuilder, motion capture, real-time character animation, MetaHuman lip sync, installation integration")

    add_section(doc, "EDUCATION | TALKS & EXHIBITIONS")
    add_label_line(doc, "MFA Computational Arts | ", "Goldsmiths, University of London | 2022-2024")
    add_label_line(doc, "BA Visual Communication Design | ", "East China University of Science and Technology | 2019-2022")
    add_label_line(doc, "Speaker, AWE USA 2025 | ", "Presented Signie")
    add_label_line(doc, "I AND AI: MIRROR | ", "Immersive Arts UK-supported pop-up exhibition and performance, Inspace, Edinburgh | 2025")
    return doc


def save(doc, filename):
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / filename
    doc.save(path)
    print(path)


def main():
    save(build_ct_en(), "Siming_Wang_Creative_Technologist_EN.docx")
    save(build_ct_cn(), "Siming_Wang_Creative_Technologist_CN.docx")
    save(build_xr_en(), "Siming_Wang_XR_Engineer_EN.docx")
    save(build_master(), "Siming_Wang_Master_Resume_EN.docx")


if __name__ == "__main__":
    main()
